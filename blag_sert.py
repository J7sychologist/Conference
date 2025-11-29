import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import re
import sys
import time
import configparser


class DocumentGenerator:
    def __init__(self, cleanup_docx=True, delay_between_files=1):
        self.cleanup_docx = cleanup_docx
        self.delay_between_files = delay_between_files

    def process_template(self, template_path, output_path, replacements):
        """Заполняет шаблон документа и сохраняет"""
        try:
            doc = Document(template_path)

            # Замена в параграфах
            for paragraph in doc.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)

            # Замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)

            # Замена в заголовках и нижних колонтитулах
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    self._replace_in_paragraph(paragraph, replacements)
                for paragraph in section.footer.paragraphs:
                    self._replace_in_paragraph(paragraph, replacements)

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Ошибка при обработке шаблона {template_path}: {e}")
            return False

    def _replace_in_paragraph(self, paragraph, replacements):
        """Заменяет текст в параграфе с сохранением форматирования"""
        # Собираем весь текст параграфа
        full_text = "".join(run.text for run in paragraph.runs)

        # Проверяем, есть ли замены в тексте
        text_changed = False
        for key, value in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
                text_changed = True

        # Если были замены, обновляем текст
        if text_changed:
            # Очищаем все runs
            for run in paragraph.runs:
                run.text = ""

            # Добавляем новый текст в первый run
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                # Если нет runs, добавляем новый
                paragraph.add_run(full_text)

    def cleanup_docx_files(self, directory):
        """Удаляет все DOCX файлы в указанной директории"""
        if not self.cleanup_docx:
            print("Удаление DOCX файлов отключено в настройках")
            return

        deleted_count = 0
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.endswith(".docx"):
                    file_path = os.path.join(root, file)
                    try:
                        os.remove(file_path)
                        deleted_count += 1
                    except Exception as e:
                        print(f"Ошибка при удалении {file}: {e}")

        if deleted_count > 0:
            print(f"Удалено {deleted_count} временных DOCX файлов")

    def generate_documents(
        self, excel_file, gratitude_template, certificate_template, output_dir
    ):
        """Генерирует все документы"""
        print("Начало генерации документов...")
        print(
            f"Настройки обработки: Удаление DOCX: {'Да' if self.cleanup_docx else 'Нет'}, Задержка: {self.delay_between_files} сек"
        )

        # Создаем папки для выходных файлов
        gratitude_dir = os.path.join(output_dir, "Благодарственные_письма")
        certificate_dir = os.path.join(output_dir, "Сертификаты")

        os.makedirs(gratitude_dir, exist_ok=True)
        os.makedirs(certificate_dir, exist_ok=True)

        # Читаем данные из Excel
        try:
            df = pd.read_excel(excel_file)
            print(f"Загружено {len(df)} записей из Excel файла")
        except Exception as e:
            print(f"Ошибка при чтении Excel файла: {e}")
            return

        gratitude_docx_files = []
        certificate_docx_files = []

        successful_gratitude = 0
        successful_certificates = 0

        for index, row in df.iterrows():
            try:
                # Извлекаем данные
                participant_name = str(row["ФИО участника"]).strip()
                report_title = str(row["Название доклада"]).strip()
                supervisor_name = str(row["ФИО руководителя"]).strip()

                print(f"Обработка: {participant_name}")

                # Генерируем благодарственное письмо
                gratitude_replacements = {
                    "{ФИО_руководителя}": supervisor_name,
                    "{ФИО_участника}": participant_name,
                    "{Название_доклада}": report_title,
                }

                # Создаем безопасное имя файла
                safe_supervisor_name = re.sub(r'[<>:"/\\|?*]', "_", supervisor_name)
                gratitude_filename = f"Благодарность_{safe_supervisor_name.replace(' ', '_')}_{index+1}.docx"
                gratitude_docx_path = os.path.join(gratitude_dir, gratitude_filename)

                if self.process_template(
                    gratitude_template, gratitude_docx_path, gratitude_replacements
                ):
                    gratitude_docx_files.append(gratitude_docx_path)
                    successful_gratitude += 1
                    print(f"  Создано благодарственное письмо")
                else:
                    print(f"  Ошибка при создании благодарственного письма")

                # Генерируем сертификат
                certificate_replacements = {
                    "{ФИО_участника}": participant_name,
                    "{Название_доклада}": report_title,
                }

                safe_participant_name = re.sub(r'[<>:"/\\|?*]', "_", participant_name)
                certificate_filename = f"Сертификат_{safe_participant_name.replace(' ', '_')}_{index+1}.docx"
                certificate_docx_path = os.path.join(
                    certificate_dir, certificate_filename
                )

                if self.process_template(
                    certificate_template,
                    certificate_docx_path,
                    certificate_replacements,
                ):
                    certificate_docx_files.append(certificate_docx_path)
                    successful_certificates += 1
                    print(f"  Создан сертификат")
                else:
                    print(f"  Ошибка при создании сертификата")

                # Задержка между обработкой участников
                if self.delay_between_files > 0 and index < len(df) - 1:
                    time.sleep(self.delay_between_files)

            except Exception as e:
                print(f"Ошибка при обработке строки {index}: {e}")
                continue

        # Конвертируем DOCX в PDF
        print("\n" + "=" * 60)
        print("Конвертация в PDF...")

        # Благодарственные письма
        print("\nКонвертация благодарственных писем:")
        for docx_file in gratitude_docx_files:
            pdf_file = docx_file.replace(".docx", ".pdf")
            try:
                convert(docx_file, pdf_file)
                print(f"  Создан PDF: {os.path.basename(pdf_file)}")

                # Задержка между конвертацией файлов
                if (
                    self.delay_between_files > 0
                    and docx_file != gratitude_docx_files[-1]
                ):
                    time.sleep(self.delay_between_files)

            except Exception as e:
                print(f"  Ошибка при конвертации: {os.path.basename(docx_file)}")

        # Сертификаты
        print("\nКонвертация сертификатов:")
        for docx_file in certificate_docx_files:
            pdf_file = docx_file.replace(".docx", ".pdf")
            try:
                convert(docx_file, pdf_file)
                print(f"  Создан PDF: {os.path.basename(pdf_file)}")

                # Задержка между конвертацией файлов
                if (
                    self.delay_between_files > 0
                    and docx_file != certificate_docx_files[-1]
                ):
                    time.sleep(self.delay_between_files)

            except Exception as e:
                print(f"  Ошибка при конвертации: {os.path.basename(docx_file)}")

        # Удаляем DOCX файлы после конвертации
        print("\nОчистка временных файлов...")
        self.cleanup_docx_files(gratitude_dir)
        self.cleanup_docx_files(certificate_dir)

        # Итоговая статистика
        print("\n" + "=" * 60)
        print("ГЕНЕРАЦИЯ ДОКУМЕНТОВ ЗАВЕРШЕНА!")
        print(f"Статистика:")
        print(f"   Благодарственные письма: {successful_gratitude}/{len(df)}")
        print(f"   Сертификаты: {successful_certificates}/{len(df)}")
        print(f"   Результаты в папке: {output_dir}")


class Config:
    """Класс для работы с конфигурационным файлом"""

    def __init__(self, config_path="config.ini"):
        self.config_path = config_path
        self.config = configparser.ConfigParser()

    def load_config(self):
        """Загружает конфигурацию из файла"""
        if not os.path.exists(self.config_path):
            print(f"Конфигурационный файл не найден: {self.config_path}")
            return False

        try:
            self.config.read(self.config_path, encoding="utf-8")
            print("Конфигурационный файл загружен")
            return True
        except Exception as e:
            print(f"Ошибка при чтении конфигурационного файла: {e}")
            return False

    def get(self, section, option, fallback=None):
        """Получает значение из конфигурации"""
        try:
            return self.config.get(section, option, fallback=fallback)
        except (configparser.NoSectionError, configparser.NoOptionError):
            return fallback

    def getboolean(self, section, option, fallback=False):
        """Получает булево значение из конфигурации"""
        try:
            return self.config.getboolean(section, option, fallback=fallback)
        except (configparser.NoSectionError, configparser.NoOptionError):
            return fallback

    def getint(self, section, option, fallback=0):
        """Получает целочисленное значение из конфигурации"""
        try:
            return self.config.getint(section, option, fallback=fallback)
        except (configparser.NoSectionError, configparser.NoOptionError):
            return fallback


def get_script_directory():
    """Возвращает путь к директории исполняемого файла или скрипта"""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))


def main():
    print("=" * 60)
    print("ГЕНЕРАТОР СЕРТИФИКАТОВ И БЛАГОДАРСТВЕННЫХ ПИСЕМ")
    print("=" * 60)

    # Получаем директорию скрипта
    script_dir = get_script_directory()
    print(f"Рабочая директория: {script_dir}")

    # Загружаем конфигурацию
    config = Config(os.path.join(script_dir, "config.ini"))
    if not config.load_config():
        input("\nНажмите Enter для выхода...")
        return

    # Получаем настройки из конфигурации
    EXCEL_FILE = os.path.join(script_dir, config.get("files", "excel_file"))
    GRATITUDE_TEMPLATE = os.path.join(
        script_dir, config.get("files", "gratitude_template")
    )
    CERTIFICATE_TEMPLATE = os.path.join(
        script_dir, config.get("files", "certificate_template")
    )
    OUTPUT_DIR = os.path.join(script_dir, config.get("paths", "output_dir"))
    CLEANUP_DOCX = config.getboolean("processing", "cleanup_docx", fallback=True)
    DELAY_BETWEEN_FILES = config.getint("processing", "delay_between_files", fallback=1)

    print("\nПоиск необходимых файлов...")

    # Проверяем существование необходимых файлов
    required_files = {
        "Excel файл": EXCEL_FILE,
        "Шаблон благодарственного письма": GRATITUDE_TEMPLATE,
        "Шаблон сертификата": CERTIFICATE_TEMPLATE,
    }

    missing_files = []
    for name, path in required_files.items():
        if os.path.exists(path):
            print(f"Найден: {name}")
        else:
            print(f"Отсутствует: {name}")
            missing_files.append(name)

    if missing_files:
        print(f"\nОшибка: Отсутствуют необходимые файлы:")
        for file in missing_files:
            print(f"   - {file}")
        print(f"\nУбедитесь, что файлы находятся в папке с программой:")
        print(f"   {script_dir}")
        input("\nНажмите Enter для выхода...")
        return

    print("\nВсе файлы найдены! Начинаем обработку...")

    time.sleep(1)

    # Создаем генератор документов с настройками из конфига
    generator = DocumentGenerator(
        cleanup_docx=CLEANUP_DOCX, delay_between_files=DELAY_BETWEEN_FILES
    )

    # Генерируем документы
    generator.generate_documents(
        excel_file=EXCEL_FILE,
        gratitude_template=GRATITUDE_TEMPLATE,
        certificate_template=CERTIFICATE_TEMPLATE,
        output_dir=OUTPUT_DIR,
    )

    print("\n" + "=" * 60)
    input("Нажмите Enter для выхода...")


if __name__ == "__main__":
    main()