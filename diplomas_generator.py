import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import re
import sys
import time
import configparser
from PyPDF2 import PdfMerger


class DiplomaGenerator:
    def __init__(self, cleanup_docx=True, delay_between_files=1):
        self.cleanup_docx = cleanup_docx
        self.delay_between_files = delay_between_files

    def load_config(self):
        """Загружает конфигурацию из config.ini"""
        config = configparser.ConfigParser()

        if getattr(sys, "frozen", False):
            script_dir = os.path.dirname(sys.executable)
        else:
            script_dir = os.path.dirname(os.path.abspath(__file__))

        config_path = os.path.join(script_dir, "config.ini")

        if not os.path.exists(config_path):
            raise FileNotFoundError(f"Файл config.ini не найден по пути: {config_path}")

        config.read(config_path, encoding="utf-8")
        return config

    def get_script_directory(self):
        """Возвращает путь к директории исполняемого файла или скрипта"""
        if getattr(sys, "frozen", False):
            return os.path.dirname(sys.executable)
        else:
            return os.path.dirname(os.path.abspath(__file__))

    def get_external_file_path(self, filename):
        """Возвращает полный путь к внешнему файлу рядом с исполняемым файлом"""
        script_dir = self.get_script_directory()
        return os.path.join(script_dir, filename)

    def replace_text_in_paragraph(self, paragraph, replacements):
        """Заменяет текст в параграфе с сохранением форматирования"""
        for search_text, replace_text in replacements.items():
            if search_text in paragraph.text:
                # Сохраняем оригинальное форматирование
                original_runs = paragraph.runs
                if not original_runs:
                    paragraph.text = paragraph.text.replace(search_text, replace_text)
                    continue

                # Собираем полный текст
                full_text = ""
                for run in original_runs:
                    full_text += run.text

                if search_text in full_text:
                    # Очищаем параграф
                    paragraph.clear()

                    # Разделяем текст на части и добавляем с форматированием
                    parts = full_text.split(search_text)

                    for i, part in enumerate(parts):
                        if part:
                            new_run = paragraph.add_run(part)
                            # Копируем форматирование из первого run
                            first_run = original_runs[0]
                            new_run.bold = first_run.bold
                            new_run.italic = first_run.italic
                            new_run.underline = first_run.underline
                            if first_run.font.size:
                                new_run.font.size = first_run.font.size
                            if first_run.font.name:
                                new_run.font.name = first_run.font.name

                        if i < len(parts) - 1:
                            new_run = paragraph.add_run(replace_text)
                            # Копируем форматирование из первого run
                            first_run = original_runs[0]
                            new_run.bold = first_run.bold
                            new_run.italic = first_run.italic
                            new_run.underline = first_run.underline
                            if first_run.font.size:
                                new_run.font.size = first_run.font.size
                            if first_run.font.name:
                                new_run.font.name = first_run.font.name

    def create_diploma_from_template(self, template_path, replacements):
        """Создает заполненный диплом на основе шаблона"""
        try:
            doc = Document(template_path)

            # Замена в параграфах
            for paragraph in doc.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)

            # Замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, replacements)

            return doc
        except Exception as e:
            print(f"[ОШИБКА] Ошибка при создании диплома: {e}")
            return None

    def merge_pdfs(self, pdf_files, output_path):
        """Объединяет несколько PDF файлов в один"""
        try:
            merger = PdfMerger()

            for pdf_file in pdf_files:
                if os.path.exists(pdf_file):
                    merger.append(pdf_file)

            merger.write(output_path)
            merger.close()
            return True
        except Exception as e:
            print(f"[ОШИБКА] Ошибка при объединении PDF файлов: {e}")
            return False

    def merge_docx_files(self, docx_files, output_path):
        """Объединяет несколько DOCX файлов в один"""
        try:
            if not docx_files:
                return False

            # Создаем новый документ для объединения
            merged_doc = Document()

            for i, docx_file in enumerate(docx_files):
                if os.path.exists(docx_file):
                    # Открываем каждый документ
                    doc = Document(docx_file)

                    # Для первого документа просто копируем содержимое
                    if i == 0:
                        for element in doc.element.body:
                            merged_doc.element.body.append(element)
                    else:
                        # Добавляем разрыв страницы перед следующим документом
                        merged_doc.add_page_break()

                        # Копируем содержимое документа
                        for element in doc.element.body:
                            merged_doc.element.body.append(element)

            merged_doc.save(output_path)
            return True

        except Exception as e:
            print(f"[ОШИБКА] Ошибка при объединении DOCX файлов: {e}")
            return False

    def cleanup_docx_files(self, directory):
        """Удаляет все DOCX файлы в указанной директории"""
        if not self.cleanup_docx:
            print("[ИНФО] Удаление DOCX файлов отключено в настройках")
            return

        deleted_count = 0
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.endswith(".docx"):
                    file_path = os.path.join(root, file)
                    try:
                        os.remove(file_path)
                        deleted_count += 1
                    except Exception as e:
                        print(f"[ОШИБКА] Ошибка при удалении {file}: {e}")

        if deleted_count > 0:
            print(f"[ИНФО] Удалено {deleted_count} временных DOCX файлов")

    def generate_diplomas(self):
        """Генерирует дипломы для призеров"""
        print("Начало генерации дипломов...")

        # Загружаем конфигурацию
        try:
            config = self.load_config()
            excel_file = self.get_external_file_path(config.get("files", "excel_file"))
            diploma_template = self.get_external_file_path(
                config.get("files", "winner_template")
            )
            output_dir = self.get_external_file_path(config.get("paths", "output_dir"))
            cleanup_docx = config.getboolean(
                "processing", "cleanup_docx", fallback=True
            )
            delay_between_files = config.getint(
                "processing", "delay_between_files", fallback=1
            )

            # Обновляем настройки из конфига
            self.cleanup_docx = cleanup_docx
            self.delay_between_files = delay_between_files

        except Exception as e:
            print(f"[ОШИБКА] Ошибка загрузки конфигурации: {e}")
            return

        print(
            f"[ИНФО] Настройки обработки: Удаление DOCX: {'Да' if self.cleanup_docx else 'Нет'}, Задержка: {self.delay_between_files} сек"
        )

        # Проверяем существование файлов
        missing_files = []
        if not os.path.exists(excel_file):
            missing_files.append(config.get("files", "excel_file"))
        if not os.path.exists(diploma_template):
            missing_files.append(config.get("files", "winner_template"))

        if missing_files:
            print("[ОШИБКА] Отсутствуют необходимые файлы:")
            for file in missing_files:
                print(f"   - {file}")
            print(
                f"Убедитесь, что файлы находятся в папке: {self.get_script_directory()}"
            )
            return

        # Создаем папку для выходных файлов
        winners_dir = os.path.join(output_dir, "Дипломы_призеров")
        os.makedirs(winners_dir, exist_ok=True)

        # Читаем данные из Excel
        try:
            df = pd.read_excel(excel_file)
            print(f"[УСПЕХ] Загружено {len(df)} записей из Excel файла")
        except Exception as e:
            print(f"[ОШИБКА] Ошибка при чтении Excel файла: {e}")
            return

        # Фильтруем призеров
        prize_winners = df[df["Призер"].isin([1, 2, 3])]
        print(f"[ИНФО] Найдено {len(prize_winners)} призеров")

        if len(prize_winners) == 0:
            print(
                "[ОШИБКА] Призеры не найдены. Проверьте столбец 'Призер' в Excel файле."
            )
            return

        successful_diplomas = 0
        individual_pdf_files = []
        individual_docx_files = []

        print("\nСоздание индивидуальных дипломов...")

        for index, row in prize_winners.iterrows():
            try:
                # Извлекаем данные
                participant_name = str(row["ФИО участника"]).strip()
                report_title = str(row["Название доклада"]).strip()
                supervisor_name = str(row["ФИО руководителя"]).strip()
                prize_level = int(row["Призер"])

                # Определяем место
                prize_text = {1: "I место", 2: "II место", 3: "III место"}.get(
                    prize_level, ""
                )

                print(f"Обрабатываем: {participant_name} ({prize_text})")

                # Подготовка замен
                replacements = {
                    "{ФИО_участника}": participant_name,
                    "{Название_доклада}": report_title,
                    "{ФИО_руководителя}": supervisor_name,
                }

                # Создаем безопасное имя файла
                safe_name = re.sub(r'[<>:"/\\|?*]', "_", participant_name)
                individual_docx_path = os.path.join(
                    winners_dir, f"Диплом_{safe_name.replace(' ', '_')}.docx"
                )
                individual_pdf_path = os.path.join(
                    winners_dir, f"Диплом_{safe_name.replace(' ', '_')}.pdf"
                )

                # Создаем индивидуальный диплом
                individual_doc = self.create_diploma_from_template(
                    diploma_template, replacements
                )

                if individual_doc:
                    # Сохраняем DOCX
                    individual_doc.save(individual_docx_path)
                    individual_docx_files.append(individual_docx_path)

                    # Конвертируем в PDF
                    try:
                        convert(individual_docx_path, individual_pdf_path)
                        individual_pdf_files.append(individual_pdf_path)
                        successful_diplomas += 1
                        print(
                            f"  [УСПЕХ] Созданы файлы: {os.path.basename(individual_docx_path)} и {os.path.basename(individual_pdf_path)}"
                        )

                    except Exception as e:
                        print(
                            f"  [ОШИБКА] Ошибка при создании PDF для {participant_name}: {e}"
                        )
                else:
                    print(
                        f"  [ОШИБКА] Ошибка при создании диплома для {participant_name}"
                    )

                # Задержка между обработкой участников
                if self.delay_between_files > 0 and index < len(prize_winners) - 1:
                    time.sleep(self.delay_between_files)

            except Exception as e:
                print(f"[ОШИБКА] Ошибка при обработке строки {index}: {e}")
                continue

        # Удаляем DOCX файлы если включено в настройках
        if self.cleanup_docx:
            print("\n🧹 Очистка временных DOCX файлов...")
            self.cleanup_docx_files(winners_dir)

        # Объединяем индивидуальные PDF файлы в один общий
        if individual_pdf_files:
            print("\nОбъединение индивидуальных PDF файлов...")
            combined_pdf_path = os.path.join(winners_dir, "Все_дипломы_призеров.pdf")

            if self.merge_pdfs(individual_pdf_files, combined_pdf_path):
                print(
                    f"  [УСПЕХ] Создан объединенный PDF: {os.path.basename(combined_pdf_path)}"
                )
                print(f"  [ИНФО] Объединено PDF файлов: {len(individual_pdf_files)}")
            else:
                print(f"  [ОШИБКА] Ошибка при создании объединенного PDF")

        # Объединяем индивидуальные DOCX файлы в один общий (только если не удалены)
        if individual_docx_files and not self.cleanup_docx:
            print("\nОбъединение индивидуальных DOCX файлов...")
            combined_docx_path = os.path.join(winners_dir, "Все_дипломы_призеров.docx")

            if self.merge_docx_files(individual_docx_files, combined_docx_path):
                print(
                    f"  [УСПЕХ] Создан объединенный DOCX: {os.path.basename(combined_docx_path)}"
                )
                print(f"  [ИНФО] Объединено DOCX файлов: {len(individual_docx_files)}")
            else:
                print(f"  [ОШИБКА] Ошибка при создании объединенного DOCX")
        elif self.cleanup_docx:
            print("\n[ИНФО] Объединение DOCX пропущено - файлы удалены по настройкам")

        # Итоговая статистика
        print("\n" + "=" * 60)
        print("ГЕНЕРАЦИЯ ДИПЛОМОВ ЗАВЕРШЕНА!")
        print(f"Статистика:")
        print(f"   Обработано призеров: {successful_diplomas}/{len(prize_winners)}")
        print(f"   Создано DOCX файлов: {len(individual_docx_files)}")
        print(f"   Создано PDF файлов: {len(individual_pdf_files)}")
        print(f"   Удаление DOCX: {'Включено' if self.cleanup_docx else 'Отключено'}")
        print(f"   Результаты в папке: {winners_dir}")
        print(f"   Объединенный PDF: Все_дипломы_призеров.pdf")
        if not self.cleanup_docx:
            print(f"   Объединенный DOCX: Все_дипломы_призеров.docx")


def main():
    print("=" * 60)
    print("ГЕНЕРАТОР ДИПЛОМОВ ДЛЯ ПРИЗЕРОВ КОНФЕРЕНЦИИ")
    print("=" * 60)

    generator = DiplomaGenerator()

    # Запускаем генерацию дипломов
    generator.generate_diplomas()

    print("\n" + "=" * 60)
    input("Нажмите Enter для выхода...")


if __name__ == "__main__":
    main()
