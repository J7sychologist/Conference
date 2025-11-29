import pandas as pd
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import time
from comtypes import client
import configparser
import sys


def get_script_directory():
    """Возвращает путь к директории исполняемого файла или скрипта"""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))


def load_config():
    """Загружает конфигурацию из config.ini рядом с исполняемым файлом"""
    config = configparser.ConfigParser()
    script_dir = get_script_directory()
    config_path = os.path.join(script_dir, "config.ini")

    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Файл config.ini не найден по пути: {config_path}")

    config.read(config_path, encoding="utf-8")
    return config


def get_external_file_path(filename):
    """Возвращает полный путь к внешнему файлу рядом с исполняемым файлом"""
    script_dir = get_script_directory()
    return os.path.join(script_dir, filename)


def check_required_files(config):
    """Проверяет наличие всех необходимых файлов"""
    script_dir = get_script_directory()
    missing_files = []

    excel_file = get_external_file_path(config.get("files", "excel_file"))
    template_file = get_external_file_path(config.get("files", "invitation_template"))
    email_template_file = get_external_file_path(config.get("files", "email_template"))
    config_file = get_external_file_path("config.ini")

    if not os.path.exists(excel_file):
        missing_files.append(config.get("files", "excel_file"))

    if not os.path.exists(template_file):
        missing_files.append(config.get("files", "invitation_template"))

    if not os.path.exists(email_template_file):
        missing_files.append(config.get("files", "email_template"))

    if not os.path.exists(config_file):
        missing_files.append("config.ini")

    return missing_files


def docx_to_pdf(docx_path, pdf_path):
    """Конвертирует DOCX в PDF используя Word"""
    try:
        word = client.CreateObject("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"Ошибка конвертации в PDF: {e}")
        return False


def replace_text_keeping_formatting(paragraph, search_text, replace_text):
    """Заменяет текст в параграфе, сохраняя форматирование"""
    if search_text in paragraph.text:
        original_runs = paragraph.runs
        if not original_runs:
            paragraph.text = paragraph.text.replace(search_text, replace_text)
            return

        full_text = ""
        for run in original_runs:
            full_text += run.text

        if search_text in full_text:
            p = paragraph._element
            p.clear()

            parts = full_text.split(search_text)

            for i, part in enumerate(parts):
                if part:
                    new_run = paragraph.add_run(part)
                    if original_runs:
                        first_run = original_runs[0]
                        new_run.bold = first_run.bold
                        new_run.italic = first_run.italic
                        new_run.underline = first_run.underline
                        if first_run.font.size:
                            new_run.font.size = first_run.font.size
                        if first_run.font.name:
                            new_run.font.name = first_run.font.name

                if i < len(parts) - 1:
                    new_run = paragraph.add_run(replace_text)
                    if original_runs:
                        first_run = original_runs[0]
                        new_run.bold = first_run.bold
                        new_run.italic = first_run.italic
                        new_run.underline = first_run.underline
                        if first_run.font.size:
                            new_run.font.size = first_run.font.size
                        if first_run.font.name:
                            new_run.font.name = first_run.font.name


def create_personalized_invitation(
    template_path, output_dir, fio, paper_title, cleanup_docx=True
):
    """Создает персонализированное приглашение в PDF"""
    try:
        doc = Document(template_path)

        # Обрабатываем все параграфы
        for paragraph in doc.paragraphs:
            # Заменяем ФИО_участника и выравниваем по центру
            if "{ФИО_участника}" in paragraph.text:
                replace_text_keeping_formatting(paragraph, "{ФИО_участника}", fio)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Заменяем Название_доклада и выравниваем по ширине
            if "{Название_доклада*}" in paragraph.text:
                replace_text_keeping_formatting(
                    paragraph, "{Название_доклада*}", paper_title
                )
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif "{Название_доклада}" in paragraph.text:
                replace_text_keeping_formatting(
                    paragraph, "{Название_доклада}", paper_title
                )
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{ФИО_участника}" in paragraph.text:
                            replace_text_keeping_formatting(
                                paragraph,
                                "{ФИО_участника}",
                                fio,
                            )
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        if "{Название_доклада*}" in paragraph.text:
                            replace_text_keeping_formatting(
                                paragraph, "{Название_доклада*}", paper_title
                            )
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        elif "{Название_доклада}" in paragraph.text:
                            replace_text_keeping_formatting(
                                paragraph, "{Название_доклада}", paper_title
                            )
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        script_dir = get_script_directory()
        output_dir_full = os.path.join(script_dir, output_dir, "Приглашения")
        os.makedirs(output_dir_full, exist_ok=True)

        temp_docx = os.path.join(
            output_dir_full, f"Приглашение_{fio.replace(' ', '_')}.docx"
        )
        doc.save(temp_docx)

        pdf_path = os.path.join(
            output_dir_full, f"Приглашение_{fio.replace(' ', '_')}.pdf"
        )

        if docx_to_pdf(os.path.abspath(temp_docx), os.path.abspath(pdf_path)):
            # Удаляем временный DOCX файл если включено в настройках
            if cleanup_docx:
                os.remove(temp_docx)
            return pdf_path
        else:
            return None

    except Exception as e:
        print(f"Ошибка создания приглашения для {fio}: {e}")
        return None


def load_email_template(config):
    """Загружает HTML шаблон письма из файла"""
    script_dir = get_script_directory()
    email_template_filename = config.get("files", "email_template")
    template_path = os.path.join(script_dir, email_template_filename)

    with open(template_path, "r", encoding="utf-8") as file:
        return file.read()


def create_email_body(fio, paper_title, config):
    """Создает тело письма из шаблона"""
    template = load_email_template(config)
    return template.replace("{fio}", fio).replace("{paper_title}", paper_title)


def send_email_simple(
    sender_email, sender_password, recipient_email, fio, paper_title, pdf_path, config
):
    """Упрощенная функция отправки письма"""
    try:
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg["Subject"] = "Приглашение на конференцию «IХ Ставеровские чтения»"

        # HTML тело письма из шаблона
        html_body = create_email_body(fio, paper_title, config)
        msg.attach(MIMEText(html_body, "html", "utf-8"))

        # Прикрепляем PDF файл
        with open(pdf_path, "rb") as file:
            attachment = MIMEApplication(file.read(), _subtype="pdf")
            attachment.add_header(
                "Content-Disposition",
                "attachment",
                filename=f"Приглашение_на_конференцию_{fio.replace(' ', '_')}.pdf",
            )
            msg.attach(attachment)

        # Получаем SMTP настройки из конфига
        smtp_server = config.get("email", "smtp_server")
        smtp_port = config.getint("email", "smtp_port")

        # Отправляем письмо
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()

        return True

    except Exception as e:
        print(f"Ошибка отправки письма для {fio}: {e}")
        return False


def wait_for_keypress():
    """Ожидает нажатия любой клавиши перед закрытием консоли"""
    print("\n" + "=" * 80)
    print("Нажмите любую клавишу для выхода...")
    try:
        import msvcrt

        msvcrt.getch()
    except ImportError:
        input()


def main():
    """Основная функция"""
    print("=" * 80)
    print("    РАССЫЛКА ПРИГЛАШЕНИЙ")
    print("=" * 80)

    try:
        # Загружаем конфигурацию
        config = load_config()

        # Получаем настройки обработки
        cleanup_docx = config.getboolean("processing", "cleanup_docx", fallback=True)
        delay_between_files = config.getint(
            "processing", "delay_between_files", fallback=2
        )

        print(
            f"Настройки обработки: Удаление DOCX: {'Да' if cleanup_docx else 'Нет'}, Задержка: {delay_between_files} сек"
        )

        # Проверяем наличие файлов
        missing_files = check_required_files(config)
        if missing_files:
            print("Отсутствуют необходимые файлы:")
            for file in missing_files:
                print(f"   - {file}")
            print(f"Убедитесь, что файлы находятся в папке: {get_script_directory()}")
            wait_for_keypress()
            return

        # Получаем полные пути к файлам
        excel_file = get_external_file_path(config.get("files", "excel_file"))
        template_file = get_external_file_path(
            config.get("files", "invitation_template")
        )
        sender_email = config.get("email", "sender_email")
        sender_password = config.get("email", "sender_password")
        output_dir = get_external_file_path(config.get("paths", "output_dir"))

        print(f"Рабочая директория: {get_script_directory()}")
        print(f"Excel файл: {excel_file}")
        print(f"Шаблон DOCX: {template_file}")

        # Читаем Excel файл
        print(f"Читаем файл: {excel_file}")
        df = pd.read_excel(excel_file)
        print(f"Найдено {len(df)} участников")

        # Проверяем колонки
        required_columns = ["ФИО участника", "Название доклада", "e-mail"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            print(f"В файле отсутствуют колонки: {missing_columns}")
            print(f"   Найдены колонки: {list(df.columns)}")
            wait_for_keypress()
            return

        # Счетчики
        pdf_created = 0
        emails_sent = 0
        emails_failed = 0
        errors = 0

        print("Начинаем обработку...")

        for index, row in df.iterrows():
            try:
                # Извлекаем данные
                fio = (
                    str(row["ФИО участника"]).strip()
                    if pd.notna(row["ФИО участника"])
                    else ""
                )
                paper_title = (
                    str(row["Название доклада"]).strip()
                    if pd.notna(row["Название доклада"])
                    else ""
                )
                email = str(row["e-mail"]).strip() if pd.notna(row["e-mail"]) else ""

                # Пропускаем пустые строки
                if not all([fio, paper_title, email]):
                    print(f"Строка {index+1}: пропущена (неполные данные)")
                    errors += 1
                    continue

                print(f"Обрабатываем: {fio}")

                # Создаем персонализированное приглашение в PDF
                pdf_path = create_personalized_invitation(
                    template_file, output_dir, fio, paper_title, cleanup_docx
                )

                if pdf_path and os.path.exists(pdf_path):
                    pdf_created += 1
                    print(f"   PDF создан: {os.path.basename(pdf_path)}")

                    # Отправляем письмо упрощенным способом
                    if send_email_simple(
                        sender_email,
                        sender_password,
                        email,
                        fio,
                        paper_title,
                        pdf_path,
                        config,
                    ):
                        emails_sent += 1
                        print(f"   Письмо отправлено: {email}")
                    else:
                        emails_failed += 1
                        print(f"   Ошибка отправки письма")

                    # Задержка между отправками
                    if delay_between_files > 0 and index < len(df) - 1:
                        time.sleep(delay_between_files)

                else:
                    errors += 1
                    print(f"   Ошибка создания PDF")

            except Exception as e:
                errors += 1
                print(f"Ошибка обработки строки {index+1}: {e}")

        # Итоги
        invitations_dir = os.path.join(output_dir, "Приглашения")
        print(f"\n{'='*80}")
        print("ИТОГИ РАССЫЛКИ:")
        print(f"   Создано PDF файлов: {pdf_created}")
        print(f"   Успешно отправлено: {emails_sent}")
        print(f"   Не отправлено: {emails_failed}")
        print(f"   Ошибок обработки: {errors}")
        print(f"   Всего участников: {len(df)}")
        print(f"   Удаление DOCX: {'Включено' if cleanup_docx else 'Отключено'}")
        print(f"   PDF файлы сохранены в: {invitations_dir}")
        print(f"{'='*80}")

    except Exception as e:
        print(f"Ошибка: {e}")

    # Ожидаем нажатия клавиши перед закрытием
    wait_for_keypress()


if __name__ == "__main__":
    main()
